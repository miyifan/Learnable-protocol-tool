# protocol_manager.py - 协议管理和存储模块
import os
import json
from pathlib import Path
import struct

class ProtocolManager:
    """协议管理类：处理协议的加载、保存和查询功能"""
    
    def __init__(self, data_dir="protocols"):
        # 确保协议存储目录存在
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(exist_ok=True)
        self.protocols = {}  # 存储所有协议
        self.protocol_commands = {}  # 存储协议指令
        self.load_all_protocols()
    
    def load_all_protocols(self):
        """加载所有协议和命令"""
        try:
            # 加载协议
            for file_path in self.data_dir.glob("**/protocol.json"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    protocol = json.load(f)
                    if protocol.get("type") == "protocol":
                        group = file_path.parent.name
                        self.protocols[f"{group}/{protocol['name']}"] = protocol
            
            # 加载标准命令文件 (ID.json)
            for file_path in self.data_dir.glob("**/*.json"):
                if file_path.name == "protocol.json":
                    continue
                
                # 检查是否是命令格式的文件名 (command_ID_name.json)
                if file_path.name.startswith("command_"):
                    parts = file_path.stem.split('_', 2)
                    if len(parts) >= 2:
                        command_id = parts[1]  # 提取ID部分
                        group = file_path.parent.name
                        print(f"发现命令格式文件: {file_path.name}, ID={command_id}, 组={group}")
                        
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                command = json.load(f)
                                
                                # 确保命令有正确的ID
                                if not command.get("protocol_id_hex"):
                                    command["protocol_id_hex"] = command_id
                                
                                # 转换旧格式到新格式
                                if group not in self.protocol_commands:
                                    self.protocol_commands[group] = {}
                                    
                                if command_id not in self.protocol_commands[group]:
                                    self.protocol_commands[group][command_id] = []
                                
                                # 添加到命令字典
                                if isinstance(command, list):
                                    self.protocol_commands[group][command_id].extend(command)
                                else:
                                    self.protocol_commands[group][command_id].append(command)
                                
                                # 添加到协议字典
                                self.protocols[f"{group}/{command_id}"] = command
                        except Exception as e:
                            print(f"加载命令文件 {file_path} 失败: {e}")
                continue
                
                group = file_path.parent.name
                command_id = file_path.stem
                
                with open(file_path, 'r', encoding='utf-8') as f:
                    commands = json.load(f)
                    
                # 转换旧格式到新格式
                if group not in self.protocol_commands:
                    self.protocol_commands[group] = {}
                    
                if command_id not in self.protocol_commands[group]:
                    self.protocol_commands[group][command_id] = []
                    
                if isinstance(commands, list):
                    self.protocol_commands[group][command_id].extend(commands)
                else:
                    self.protocol_commands[group][command_id].append(commands)
                    
            print(f"加载完成，协议数量: {len(self.protocols)}")
            print(f"command_numbers: {len(self.protocol_commands)}")
            return True, "协议和命令加载成功"
                except Exception as e:
            return False, f"加载协议和命令失败: {str(e)}"
    
    def save_protocol(self, protocol_data):
        """保存协议数据到文件"""
        # 确保协议数据包含十进制和十六进制形式
        # 处理新的键名：hex_id和dec_id
        if "hex_id" in protocol_data:
            protocol_data["protocol_id_hex"] = protocol_data["hex_id"]
            try:
                if "dec_id" in protocol_data:
                    protocol_data["protocol_id_dec"] = str(protocol_data["dec_id"])
                else:
                    protocol_data["protocol_id_dec"] = str(int(protocol_data["hex_id"], 16))
            except ValueError:
                protocol_data["protocol_id_dec"] = "未知"
        elif "protocol_id_hex" not in protocol_data and "protocol_id" in protocol_data:
            protocol_data["protocol_id_hex"] = protocol_data["protocol_id"]
            try:
                protocol_data["protocol_id_dec"] = str(int(protocol_data["protocol_id"], 16))
            except ValueError:
                protocol_data["protocol_id_dec"] = "未知"
        
        # 获取协议ID
        protocol_id = protocol_data.get("protocol_id_hex", "unknown")
        protocol_name = protocol_data.get("name", "").lower()
        protocol_type = protocol_data.get("type", "protocol")
        
        # 打印保存信息
        print(f"正在保存{'协议' if protocol_type == 'protocol' else '命令'}: {protocol_name} (ID: {protocol_id})")
        if 'fields' in protocol_data:
            print(f"字段数量: {len(protocol_data['fields'])}")
            for i, field in enumerate(protocol_data['fields']):
                print(f"  字段 {i+1}: {field.get('name', '')} ({field.get('type', '')}) 位置:{field.get('start_pos', 0)}-{field.get('end_pos', 0)}")
        
        # 确定存储目录
        if protocol_type == "protocol":
            # 协议直接存储在protocols目录下的协议名子目录中
            group = protocol_name
            protocol_data["group"] = group
            # 创建协议目录
            protocol_dir = self.data_dir / group
            protocol_dir.mkdir(exist_ok=True, parents=True)
            # 协议文件保存为protocol.json
            file_path = protocol_dir / "protocol.json"
        else:  # command类型
            # 命令存储在协议名目录下
            parent_protocol_name = protocol_data.get("protocol_name", "")
            if not parent_protocol_name:
                return False, "命令必须关联到一个协议"
            
            group = protocol_data.get("group", "")
            if not group:
                group = parent_protocol_name.lower()
                protocol_data["group"] = group
            
            # 创建协议命令目录
            protocol_dir = self.data_dir / group
            protocol_dir.mkdir(exist_ok=True, parents=True)
            # 命令文件以命令ID命名
            file_path = protocol_dir / f"{protocol_id}.json"
        
        try:
            print(f"保存到文件: {file_path}")
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(protocol_data, f, ensure_ascii=False, indent=2)
            
            # 更新内存中的协议数据
            full_key = f"{group}/{protocol_id}" if group else protocol_id
            
            # 根据类型更新不同的数据结构
            if protocol_type == "protocol":
                self.protocols[full_key] = protocol_data
            elif protocol_type == "command":
                protocol_name = protocol_data.get("protocol_name", "")
                command_name = protocol_data.get("name", "")
                
                print(f"保存命令到协议: {protocol_name}, 命令ID: {protocol_id}, 名称: {command_name}")
                
                # 确保protocol_commands中有对应的协议名和命令ID
                if protocol_name not in self.protocol_commands:
                    self.protocol_commands[protocol_name] = {}
                
                if protocol_id not in self.protocol_commands[protocol_name]:
                    self.protocol_commands[protocol_name][protocol_id] = []
                
                # 检查是否已存在相同名称的命令
                command_exists = False
                commands_list = self.protocol_commands[protocol_name][protocol_id]
                
                # 如果命令列表不是列表类型，转换为列表
                if not isinstance(commands_list, list):
                    if isinstance(commands_list, dict):
                        commands_list = [commands_list]
                    else:
                        commands_list = []
                    self.protocol_commands[protocol_name][protocol_id] = commands_list
                
                # 检查是否已存在同名命令，如果存在则更新它
                for i, cmd in enumerate(commands_list):
                    if isinstance(cmd, dict) and cmd.get("name") == command_name:
                        print(f"更新已存在的命令: {command_name}")
                        commands_list[i] = protocol_data
                        command_exists = True
                        break
                
                # 如果不存在同名命令，添加到列表中
                if not command_exists:
                    print(f"添加新命令: {command_name}")
                    commands_list.append(protocol_data)
                
                # 在protocols字典中也保存一份
                self.protocols[full_key] = protocol_data
            
            print(f"保存成功, 协议键: {full_key}")
            return True, f"{'协议' if protocol_type == 'protocol' else '命令'}已保存: {protocol_id} (十进制: {protocol_data.get('protocol_id_dec', '未知')}) 到 {group}"
        except Exception as e:
            print(f"保存失败: {e}")
            return False, f"保存{'协议' if protocol_type == 'protocol' else '命令'}失败: {e}"
    
    def delete_protocol(self, protocol_key):
        """删除指定的协议"""
        print(f"尝试删除协议，键值: {protocol_key}")
        print(f"当前协议键列表: {list(self.protocols.keys())}")
        
        if protocol_key not in self.protocols:
            print(f"检查是否是旧格式命令文件: {protocol_key}")
            # 检查是否是旧格式的命令文件名导致的问题
            if '/' in protocol_key:
                group, protocol_id = protocol_key.split('/', 1)
                print(f"分解键值: 组={group}, ID={protocol_id}")
                
                # 检查该组下的目录是否存在
                protocol_dir = self.data_dir / group
                if protocol_dir.exists():
                    # 检查是否有命令格式的文件 (command_ID_name.json)
                    for file_path in protocol_dir.glob(f"command_{protocol_id}_*.json"):
                        print(f"找到匹配文件: {file_path}")
                        try:
                            # 删除文件
                            file_path.unlink()
                            print(f"已删除文件: {file_path}")
                            
                            # 更新内存中的数据
                            # 由于是旧格式文件，可能只更新了文件但没有更新内存结构，先检查命令字典
                            protocol_name = group  # 假设协议名就是组名
                            if protocol_name in self.protocol_commands and protocol_id in self.protocol_commands[protocol_name]:
                                del self.protocol_commands[protocol_name][protocol_id]
                                print(f"从protocol_commands中删除了命令: {protocol_name}/{protocol_id}")
                            
                            return True, f"命令文件 {file_path.name} 已删除"
                        except Exception as e:
                            return False, f"删除命令文件失败: {e}"
            
            # 如果上面的步骤都没有找到匹配的文件，返回原始错误
            return False, f"协议不存在: {protocol_key}"
            
        protocol_data = self.protocols[protocol_key]
        group = protocol_data.get("group", "")
        protocol_id = protocol_data.get("protocol_id_hex", "")
        protocol_type = protocol_data.get("type", "protocol")
        protocol_name = protocol_data.get("name", "")
        
        try:
            # 确定文件路径
            if protocol_type == "protocol":
                # 如果是协议，删除protocol.json文件
                file_path = self.data_dir / group / "protocol.json"
                
                if file_path.exists():
                    file_path.unlink()
                    
                # 检查是否需要删除协议目录（如果目录为空）
                protocol_dir = self.data_dir / group
                if protocol_dir.exists():
                    # 检查目录中是否还有其他文件
                    remaining_files = list(protocol_dir.glob("*.json"))
                    if not remaining_files:
                        # 如果没有剩余文件，删除目录
                        protocol_dir.rmdir()
                        print(f"已删除空目录: {group}")
                
                # 删除该协议下的所有命令
                if protocol_name in self.protocol_commands:
                    # 先记录要删除的命令键
                    command_keys = []
                    for cmd_id, cmd in self.protocol_commands[protocol_name].items():
                        cmd_group = cmd.get("group", "")
                        cmd_key = f"{cmd_group}/{cmd_id}" if cmd_group else cmd_id
                        command_keys.append(cmd_key)
                    
                    # 从内存结构中删除命令
                    for cmd_key in command_keys:
                        if cmd_key in self.protocols:
                            del self.protocols[cmd_key]
                    
                    # 从命令字典中删除该协议的所有命令
                    del self.protocol_commands[protocol_name]
            else:
                # 如果是命令，需要检查并删除命令文件
                # 首先尝试使用标准命令文件名 (ID.json)
                standard_file_path = self.data_dir / group / f"{protocol_id}.json"
                
                # 然后尝试使用命令格式文件名 (command_ID_name.json)
                command_file_found = False
                # 如果标准文件存在，则删除
                if standard_file_path.exists():
                    standard_file_path.unlink()
                    command_file_found = True
                else:
                    # 尝试查找匹配的命令文件
                    for file_path in (self.data_dir / group).glob(f"command_{protocol_id}_*.json"):
                    file_path.unlink()
                        command_file_found = True
                        print(f"已删除命令文件: {file_path}")
                        break
                
                if not command_file_found:
                    print(f"警告: 未找到对应的命令文件")
                
                # 从命令字典中删除该命令
                protocol_name = protocol_data.get("protocol_name", "")
                if protocol_name in self.protocol_commands and protocol_id in self.protocol_commands[protocol_name]:
                    del self.protocol_commands[protocol_name][protocol_id]
            
            # 从协议字典中删除
            if protocol_key in self.protocols:
                del self.protocols[protocol_key]
                
            return True, f"{'协议' if protocol_type == 'protocol' else '命令'} {protocol_id} 已删除"
        except Exception as e:
            return False, f"删除{'协议' if protocol_type == 'protocol' else '命令'}失败: {e}"
    
    def get_protocols(self):
        """获取所有协议的列表"""
        return list(self.protocols.values())
    
    def get_protocol_by_key(self, protocol_key):
        """根据键获取协议"""
        print(f"尝试通过键获取协议: {protocol_key}")
        print(f"当前协议键列表: {list(self.protocols.keys())}")
        
        # 先从协议字典中查找
        if protocol_key in self.protocols:
            print(f"在protocols字典中找到键: {protocol_key}")
            protocol = self.protocols[protocol_key]
            if isinstance(protocol, dict):
                return protocol
            elif isinstance(protocol, list) and protocol:
                # 如果是命令列表，返回第一个命令
                return protocol[0]
            
        # 如果键包含斜杠，尝试分离组名和ID
        if '/' in protocol_key:
            group, id_part = protocol_key.split('/', 1)
            
            # 检查是否在该组下找到对应ID的命令
            if group in self.protocol_commands:
                print(f"检查组 {group} 下是否有ID为 {id_part} 的命令")
                for protocol_name, commands in self.protocol_commands.items():
                    if id_part in commands:
                        print(f"在命令字典中找到匹配: {protocol_name}/{id_part}")
                        # 返回该命令ID下的第一个命令
                        if isinstance(commands[id_part], list) and commands[id_part]:
                            return commands[id_part][0]
        
        # 直接在命令字典中查找ID
        for protocol_name, commands in self.protocol_commands.items():
            if protocol_key in commands:
                print(f"在命令字典中找到键: {protocol_key}")
                # 返回该命令ID下的第一个命令
                if isinstance(commands[protocol_key], list) and commands[protocol_key]:
                    return commands[protocol_key][0]
        
        # 如果找不到，循环打印所有协议命令键值，帮助调试
        print("未找到协议，打印当前所有命令情况:")
        for protocol_name, commands in self.protocol_commands.items():
            print(f"- 协议 {protocol_name} 的命令列表:")
            for cmd_id, cmd_list in commands.items():
                print(f"  - 命令ID: {cmd_id}")
                if isinstance(cmd_list, list):
                    for cmd in cmd_list:
                        if isinstance(cmd, dict):
                            print(f"    - 名称: {cmd.get('name', '')}")
                else:
                    print(f"    - 命令格式错误: {type(cmd_list)}")
        
        return None
    
    def find_matching_protocol(self, hex_data):
        """根据16进制数据查找匹配的协议或命令"""
        if not hex_data:
            print("未提供数据，无法查找匹配协议")
            return None
            
        print(f"查找匹配的协议/命令，原始数据: {hex_data[:20]}...")
        
        # 提取协议ID (前两个字节)
        protocol_id = hex_data[:2].upper() if len(hex_data) >= 2 else ""
        print(f"提取的协议ID(前两个字节): {protocol_id}")
        
        # 打印所有文件目录结构用于调试
        print("协议和命令文件结构:")
        self._print_protocols_directory_structure()
        
        # 首先尝试对命令进行匹配
        # 1. 直接查找A0命令（针对用户提到的情况）
        print(f"尝试直接匹配命令ID: {protocol_id}")
        # 遍历所有协议组
        for group_name, group_commands in self.protocol_commands.items():
            print(f"检查协议组: {group_name}")
            # 检查该组中是否有匹配的命令ID
            if protocol_id in group_commands:
                print(f"在组 {group_name} 中找到匹配的命令ID: {protocol_id}")
                commands = group_commands[protocol_id]
                # 返回第一个匹配的命令
                if isinstance(commands, list) and commands:
                    print(f"返回命令: {commands[0].get('name', '')}")
                    return commands[0]
                elif isinstance(commands, dict):
                    print(f"返回命令: {commands.get('name', '')}")
                    return commands
        
        # 2. 从第4个字节提取命令ID (索引6-7)并查找
        command_id = hex_data[6:8].upper() if len(hex_data) >= 8 else ""
        if command_id:
            print(f"提取的命令ID(第4字节): {command_id}")
            for group_name, group_commands in self.protocol_commands.items():
                print(f"检查协议组: {group_name}")
                if command_id in group_commands:
                    print(f"在组 {group_name} 中找到匹配的命令ID: {command_id}")
                    commands = group_commands[command_id]
                    # 返回第一个匹配的命令
                    if isinstance(commands, list) and commands:
                        print(f"返回命令: {commands[0].get('name', '')}")
                        return commands[0]
                    elif isinstance(commands, dict):
                        print(f"返回命令: {commands.get('name', '')}")
                        return commands
        
        # 3. 在protocols字典中查找
        print("在protocols字典中查找命令...")
        for protocol_key, protocol in self.protocols.items():
            if protocol.get('type') == 'command':
                protocol_id_hex = protocol.get('protocol_id_hex', '').upper()
                if protocol_id_hex == protocol_id or protocol_id_hex == command_id:
                    print(f"在protocols中找到匹配的命令: {protocol.get('name', '')}")
                    return protocol
        
        # 如果没有找到匹配的命令，尝试匹配协议
        print("尝试匹配协议...")
        for protocol_key, protocol in self.protocols.items():
            if protocol.get('type') == 'protocol':
                protocol_id_hex = protocol.get('protocol_id_hex', '').upper()
                if protocol_id_hex == protocol_id:
                    print(f"找到匹配的协议: {protocol.get('name', '')}")
                return protocol
                
        print(f"未找到匹配的协议或命令, 协议ID: {protocol_id}, 命令ID: {command_id}")
        return None
        
    def _print_protocols_directory_structure(self):
        """打印协议和命令的目录结构，用于调试"""
        try:
            import os
            print("文件系统中的协议结构:")
            protocols_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'protocols')
            if os.path.exists(protocols_dir):
                for root, dirs, files in os.walk(protocols_dir):
                    rel_path = os.path.relpath(root, protocols_dir)
                    if rel_path == '.':
                        print(f"|- {os.path.basename(root)}")
                    else:
                        print(f"|  |- {rel_path}")
                    for file in files:
                        print(f"|     |- {file}")
        except Exception as e:
            print(f"打印目录结构时出错: {e}")
            
        print("\nprotocol_commands字典结构:")
        for protocol_name, commands in self.protocol_commands.items():
            print(f"|- 协议: {protocol_name}")
            for cmd_id, cmd_list in commands.items():
                print(f"|  |- 命令ID: {cmd_id}")
                if isinstance(cmd_list, list):
                    for i, cmd in enumerate(cmd_list):
                        print(f"|     |- [{i}] {cmd.get('name', 'unknown')}")
                else:
                    print(f"|     |- (非列表) {cmd_list.get('name', 'unknown') if isinstance(cmd_list, dict) else type(cmd_list)}")

    def get_protocol_commands(self, protocol_name):
        """获取指定协议下的所有命令"""
        print(f"获取协议 '{protocol_name}' 下的所有命令")
        commands = []
        
        # 用于跟踪已添加的命令ID，避免重复
        added_command_ids = set()
        
        # 首先检查protocol_commands字典
        if protocol_name in self.protocol_commands:
            print(f"在protocol_commands中找到协议: {protocol_name}")
            for command_id, command_list in self.protocol_commands[protocol_name].items():
                print(f"处理命令ID: {command_id}")
                
                # 根据命令列表类型处理
                if isinstance(command_list, list):
                    print(f"命令是列表，包含 {len(command_list)} 个命令")
                    for cmd in command_list:
                        if isinstance(cmd, dict):
                            cmd_id = cmd.get('protocol_id_hex', '')
                            cmd_name = cmd.get('name', '')
                            cmd_key = f"{cmd_name}_{cmd_id}"
                            
                            if cmd_key not in added_command_ids:
                                commands.append(cmd)
                                added_command_ids.add(cmd_key)
                                print(f"添加命令: {cmd_name} (ID: {cmd_id})")
                            else:
                                print(f"跳过重复命令: {cmd_name} (ID: {cmd_id})")
                        else:
                            print(f"忽略非字典命令: {type(cmd)}")
                elif isinstance(command_list, dict):
                    print(f"命令是字典，添加单个命令: {command_list.get('name', 'unnamed')}")
                    cmd_id = command_list.get('protocol_id_hex', '')
                    cmd_name = command_list.get('name', '')
                    cmd_key = f"{cmd_name}_{cmd_id}"
                    
                    if cmd_key not in added_command_ids:
                        commands.append(command_list)
                        added_command_ids.add(cmd_key)
                    else:
                        print(f"跳过重复命令: {cmd_name} (ID: {cmd_id})")
                else:
                    print(f"未知命令类型: {type(command_list)}")
        else:
            print(f"在protocol_commands中未找到协议: {protocol_name}")
        
        # 然后检查protocols字典中的命令
        command_count = 0
        for protocol_key, protocol in self.protocols.items():
            if (protocol.get('type') == 'command' and 
                protocol.get('protocol_name') == protocol_name):
                cmd_id = protocol.get('protocol_id_hex', '')
                cmd_name = protocol.get('name', '')
                cmd_key = f"{cmd_name}_{cmd_id}"
                
                if cmd_key not in added_command_ids:
                    commands.append(protocol)
                    added_command_ids.add(cmd_key)
                    command_count += 1
                    print(f"在protocols中找到命令: {cmd_name} (ID: {cmd_id})")
                else:
                    print(f"在protocols中跳过重复命令: {cmd_name} (ID: {cmd_id})")
        
        print(f"在protocols中找到 {command_count} 个命令")
        print(f"总共找到 {len(commands)} 个命令")
        
        for i, cmd in enumerate(commands):
            print(f"命令 {i+1}: {cmd.get('name', 'unnamed')} (ID: {cmd.get('protocol_id_hex', 'unknown')})")
        
        return commands
    
    def get_protocol_enum(self):
        """获取所有协议的枚举值"""
        result = {}
        
        # 添加协议
        for key, protocol in self.protocols.items():
            if isinstance(protocol, dict):  # 确保是协议对象
            name = protocol.get('name', key)
            group = protocol.get('group', '')
            if group:
                name = f"[{group}] {name}"
            result[key] = name
            
        # 添加命令
        for protocol_name, commands in self.protocol_commands.items():
            for command_id, command_list in commands.items():
                if isinstance(command_list, list):  # 确保是命令列表
                    for command in command_list:
                        if isinstance(command, dict):  # 确保是命令对象
                            name = command.get('name', command_id)
                group = command.get('group', '')
                if group:
                    name = f"[{group}] {name}"
                            result[f"{command_id}"] = name
            
        return result

    def get_protocol_tree(self):
        """获取分层的协议树结构"""
        tree = {}
        
        for key, protocol in self.protocols.items():
            group = protocol.get('group', '')
            if group not in tree:
                tree[group] = []
            tree[group].append(protocol)
            
        return tree

    def add_protocol_field(self, protocol_key, field_name, field_type, start_pos, field_length):
        """添加协议字段"""
        protocol = self.get_protocol_by_key(protocol_key)
        if not protocol:
            print(f"要添加字段的协议不存在: {protocol_key}")
            return False, f"字段添加失败: 协议 {protocol_key} 不存在"
        
        print(f"向协议 {protocol_key} 添加字段: {field_name}")
        
        # 确定正确的目标对象
        target_obj = protocol
        
        # 如果是命令类型，确保目标是该命令对象
        if protocol.get('type') == 'command':
            group = protocol.get('group', '')
            command_id = protocol.get('protocol_id_hex', '')
            
            print(f"这是一个命令，组名:{group}, 命令ID:{command_id}")
            
            # 检查命令字典中是否已存在该命令
            protocol_name = protocol.get('protocol_name', '')
            if protocol_name in self.protocol_commands and command_id in self.protocol_commands[protocol_name]:
                target_obj = self.protocol_commands[protocol_name][command_id]
                print(f"找到命令对象: {target_obj.get('name', '未命名')}")
        
        # 初始化fields字段
        if 'fields' not in target_obj:
            target_obj['fields'] = []
        
        # 确保字段名不重复
        for field in target_obj['fields']:
            if field['name'] == field_name:
                return False, f"字段添加失败: 字段名 {field_name} 已存在"
        
        # 计算结束位置
        end_pos = start_pos + field_length - 1
        
        # 添加新字段
        new_field = {
            'name': field_name,
            'type': field_type,
            'start_pos': start_pos,
            'end_pos': end_pos,
            'endian': 'little'  # 默认使用小端序
        }
        target_obj['fields'].append(new_field)
        
        # 保存更新后的协议
        if protocol.get('type') == 'command':
            group = protocol.get('group', '')
            command_id = protocol.get('protocol_id_hex', '')
            success, message = self.save_command(group, command_id, target_obj)
            print(f"保存命令结果: {success}, {message}")
        else:
            success, message = self.save_protocol(target_obj)
            print(f"保存协议结果: {success}, {message}")
        
        if not success:
            return False, f"字段添加失败: {message}"
        
        return True, "字段添加成功"

    def remove_protocol_field(self, protocol_key, field_index):
        """从指定协议中删除字段"""
        if protocol_key in self.protocols:
            protocol = self.protocols[protocol_key]
            
            if 'fields' in protocol and 0 <= field_index < len(protocol['fields']):
                del protocol['fields'][field_index]
                return self.save_protocol(protocol)
        
        return False, "无法删除字段"

    def get_field_type_by_size(self, byte_count):
        """根据字节数自动判断字段类型"""
        if byte_count == 1:
            return "u8"  # 8位无符号整数
        elif byte_count == 2:
            return "u16"  # 16位无符号整数
        elif byte_count == 4:
            return "u32"  # 32位无符号整数或浮点数
        elif byte_count == 8:
            return "u64"  # 64位无符号整数或双精度浮点数
        elif byte_count > 8:
            # 对于较长的数据推荐ASCII或UTF8
            if 8 < byte_count <= 32:  # 字符串可能性较大
                return "ascii"
            else:
                return "bytes"  # 二进制数据
        else:
            return "unknown"
    
    def get_supported_field_types(self):
        """获取支持的字段类型列表"""
        return [
            # 将uxx、float、char、double、时间戳、bool放在最上面
            "u8", "u16", "u32", "u64", 
            "float", "double",
            "char", # 替代string类型
            "timestamp", 
            "bool",
            # 文本类型
            "ascii", "utf8", 
            # 特殊类型
            "date", "hex", 
            # 其他类型
            "bytes"
        ]

    def parse_protocol_data(self, hex_data, protocol):
        """解析协议数据，返回字段值"""
        if not protocol or 'fields' not in protocol:
            return None
            
        result = {
            'protocol_name': protocol.get('name', ''),
            'protocol_id': protocol.get('protocol_id_dec', ''),
            'fields': []
        }
        
        for field in protocol['fields']:
            field_result = self._parse_field(field, hex_data)
            if field_result:
                result['fields'].append(field_result)
        
        return result
    
    def _parse_field(self, field, hex_data):
        """解析单个字段"""
        try:
            start_pos = field.get('start_pos', 0)
            end_pos = field.get('end_pos', 0)
            field_type = field.get('type', 'u8')
            endian = field.get('endian', 'little')
            
            # 计算字节位置
            start_byte = start_pos * 2  # 每个字节2个16进制字符
            end_byte = (end_pos + 1) * 2
            
            # 边界检查
            if start_byte >= len(hex_data) or end_byte > len(hex_data) or start_byte >= end_byte:
                print(f"字段位置超出范围: {field.get('name', '')}，位置: {start_pos}-{end_pos}，数据长度: {len(hex_data)//2}")
                return None
            
            # 获取字段的16进制数据
            field_hex = hex_data[start_byte:end_byte]
            if not field_hex:
                return None
            
            # 根据字段类型解析值
            value = self._convert_field_value(field_hex, field_type, endian)
            
            return {
                'name': field.get('name', ''),
                'type': field_type,
                'value': value,
                'hex': field_hex,
                'description': field.get('description', ''),
                'start_pos': start_pos,
                'end_pos': end_pos
            }
        except Exception as e:
            print(f"解析字段失败: {e}, 字段: {field.get('name', '')}, 位置: {field.get('start_pos', 0)}-{field.get('end_pos', 0)}")
            return None
    
    def _convert_field_value(self, hex_data, field_type, endian='little'):
        """转换字段值为对应类型"""
        try:
            if field_type in ['u8', 'i8', 'BYTE']:
                value = int(hex_data, 16)
                if field_type == 'i8' and value > 127:
                    value -= 256
                return value
                
            elif field_type in ['u16', 'i16', 'WORD']:
                if endian == 'little':
                    value = int(hex_data[2:4] + hex_data[0:2], 16)
                else:
                    value = int(hex_data, 16)
                if field_type == 'i16' and value > 32767:
                    value -= 65536
                return value
                
            elif field_type in ['u32', 'i32', 'DWORD']:
                if endian == 'little':
                    value = int(hex_data[6:8] + hex_data[4:6] + hex_data[2:4] + hex_data[0:2], 16)
                else:
                    value = int(hex_data, 16)
                if field_type == 'i32' and value > 2147483647:
                    value -= 4294967296
                return value
                
            elif field_type in ['u64', 'i64', 'QWORD']:
                if endian == 'little':
                    value = int(hex_data[14:16] + hex_data[12:14] + hex_data[10:12] + hex_data[8:10] +
                              hex_data[6:8] + hex_data[4:6] + hex_data[2:4] + hex_data[0:2], 16)
                else:
                    value = int(hex_data, 16)
                return value
                
            elif field_type in ['float']:
                if endian == 'little':
                    hex_bytes = bytes.fromhex(hex_data[6:8] + hex_data[4:6] + hex_data[2:4] + hex_data[0:2])
                else:
                    hex_bytes = bytes.fromhex(hex_data)
                return struct.unpack('f', hex_bytes)[0]
                
            elif field_type in ['double']:
                if endian == 'little':
                    hex_bytes = bytes.fromhex(hex_data[14:16] + hex_data[12:14] + hex_data[10:12] + hex_data[8:10] +
                                            hex_data[6:8] + hex_data[4:6] + hex_data[2:4] + hex_data[0:2])
                else:
                    hex_bytes = bytes.fromhex(hex_data)
                return struct.unpack('d', hex_bytes)[0]
            
            elif field_type == 'ascii':
                # 将16进制转换为ASCII字符串，忽略不可打印字符
                try:
                    return bytes.fromhex(hex_data).decode('ascii', errors='replace')
                except:
                    return hex_data
            
            elif field_type == 'utf8':
                # 将16进制转换为UTF-8字符串
                try:
                    return bytes.fromhex(hex_data).decode('utf-8', errors='replace')
                except:
                    return hex_data
            
            elif field_type == 'char':
                # 专门用于显示字符串，优先使用UTF-8编码支持中文
                try:
                    # 检查是否只包含数字和常见ASCII字符
                    is_simple_ascii = all(c in '0123456789abcdefABCDEF' for c in hex_data)
                    # 如果全是数字，可能是纯数字字符串，直接显示原始16进制
                    if is_simple_ascii and len(hex_data) <= 8:
                        # 可能是纯数字，尝试显示数值
                        value = int(hex_data, 16)
                        return str(value)
                        
                    # 对于ASCII范围内的字符，直接用ASCII解码可能更好
                    is_ascii_range = True
                    for i in range(0, len(hex_data), 2):
                        if i+1 < len(hex_data):
                            byte_val = int(hex_data[i:i+2], 16)
                            if byte_val > 127:
                                is_ascii_range = False
                                break
                    
                    if is_ascii_range:
                        result = bytes.fromhex(hex_data).decode('ascii', errors='replace')
                        # 检查结果是否包含问号(解码失败标志)
                        if '' not in result:
                            return result
                    
                    # 尝试其他编码
                    # 先尝试UTF-8
                    utf8_result = bytes.fromhex(hex_data).decode('utf-8', errors='replace')
                    if '' not in utf8_result:
                        return utf8_result
                    
                    # 再尝试GB2312
                    try:
                        gb_result = bytes.fromhex(hex_data).decode('gb2312', errors='replace')
                        if '' not in gb_result:
                            return gb_result
                    except:
                        pass
                        
                    # 最后尝试latin1（总能成功，但可能不是正确的编码）
                    latin_result = bytes.fromhex(hex_data).decode('latin1', errors='replace')
                    
                    # 如果最终结果仍然包含大量替换字符，直接返回16进制
                    if latin_result.count('') > len(latin_result) / 3:
                        return f"0x{hex_data}"
                    
                    return latin_result
                except Exception as e:
                    print(f"字符解码失败: {e}")
                    return f"0x{hex_data}"
            
            elif field_type == 'hex':
                # 保持原始16进制形式，但格式化为带0x前缀的形式
                return '0x' + hex_data.upper()
            
            elif field_type == 'date':
                # 假设格式为: YYYYMMDD (4字节)
                if len(hex_data) == 8:
                    try:
                        # 处理小端/大端
                        if endian == 'little':
                            # 处理小端序: 如 20220314 实际存储为 14 03 22 20
                            year = int(hex_data[6:8] + hex_data[4:6], 16)
                            month = int(hex_data[2:4], 16)
                            day = int(hex_data[0:2], 16)
                        else:
                            # 大端序: 20 22 03 14
                            year = int(hex_data[0:4], 16)
                            month = int(hex_data[4:6], 16)
                            day = int(hex_data[6:8], 16)
                        
                        # 验证日期合法性
                        if 1 <= month <= 12 and 1 <= day <= 31:
                            return f"{year:04d}-{month:02d}-{day:02d}"
                    except:
                        pass
                return f"日期格式错误: {hex_data}"
            
            elif field_type == 'timestamp':
                # Unix时间戳 (4字节)
                if len(hex_data) == 8:
                    try:
                        if endian == 'little':
                            timestamp = int(hex_data[6:8] + hex_data[4:6] + hex_data[2:4] + hex_data[0:2], 16)
                        else:
                            timestamp = int(hex_data, 16)
                        
                        from datetime import datetime
                        dt = datetime.fromtimestamp(timestamp)
                        return dt.strftime("%Y-%m-%d %H:%M:%S")
                    except:
                        pass
                return f"时间戳格式错误: {hex_data}"
                
            elif field_type in ['string', 'STRING']:
                # 尝试将16进制转换为UTF-8字符串，支持中文
                try:
                    return bytes.fromhex(hex_data).decode('utf-8', errors='replace')
                except:
                    try:
                        return bytes.fromhex(hex_data).decode('gb2312', errors='replace')
                    except:
                        return hex_data
                    
            elif field_type in ['bytes', 'CUSTOM']:
                return hex_data
                
            elif field_type == 'bool':
                return bool(int(hex_data, 16))
                
            return hex_data  # 默认返回16进制字符串
            
        except Exception as e:
            print(f"转换字段值失败: {e}")
            return hex_data  # 转换失败时返回原始16进制字符串

    # 协议头相关方法
    def set_protocol_header(self, protocol_key, header_fields):
        """设置协议头字段定义"""
        if protocol_key in self.protocols:
            protocol = self.protocols[protocol_key]
            
            # 设置协议头字段
            protocol['header_fields'] = header_fields
            
            # 保存更新后的协议
            return self.save_protocol(protocol)
        
        return False, f"协议 {protocol_key} 不存在"
    
    def get_protocol_header(self, protocol_key):
        """获取协议头定义"""
        if protocol_key in self.protocols:
            protocol = self.protocols[protocol_key]
            return protocol.get('header_fields', [])
        return []

    # 文档生成相关方法
    def generate_protocol_doc(self, protocol_key=None, output_format="docx"):
        """生成协议文档，可以是单个协议或所有协议"""
        try:
            # 检查是否有协议可用
            if not self.protocols:
                return False, "没有可用的协议，请先添加协议"
                
            # 导入文档生成库
            if output_format == "docx":
                try:
                from docx import Document
                    from docx.shared import Pt, Cm, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.oxml.ns import qn
                    import datetime
                    
                document = Document()
                    
                    # 设置默认字体为微软雅黑
                    style = document.styles['Normal']
                    style.font.name = '微软雅黑'
                    style.font.size = Pt(10)
                    # 中文字体设置
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                # 设置文档标题
                title = document.add_heading('协议文档', level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 设置标题字体为微软雅黑
                    for run in title.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                # 生成协议头部分
                self._generate_header_doc(document, protocol_key)
                
                # 生成协议号列表
                self._generate_protocol_list_doc(document, protocol_key)
                
                # 生成每个协议的命令字段详情
                self._generate_protocol_fields_doc(document, protocol_key)
                
                # 保存文档
                filename = f"协议文档_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                document.save(filename)
                return True, f"文档已生成: {filename}"
                    
                except ImportError:
                    return False, "缺少python-docx库，请安装: pip install python-docx"
                except Exception as e:
                    return False, f"生成Word文档时出错: {str(e)}"
                
            elif output_format == "xlsx":
                try:
                import pandas as pd
                import openpyxl
                    from openpyxl.styles import Alignment, Font, PatternFill
                    import datetime
                
                # 创建Excel文件
                    filename = f"协议文档_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                    writer = pd.ExcelWriter(filename, engine='openpyxl')
                
                # 生成协议头表格
                self._generate_header_excel(writer, protocol_key)
                
                # 生成协议号列表
                self._generate_protocol_list_excel(writer, protocol_key)
                
                # 生成每个协议的命令字段详情
                self._generate_protocol_fields_excel(writer, protocol_key)
                
                    # 设置所有工作表的字体为微软雅黑
                    workbook = writer.book
                    for sheet_name in workbook.sheetnames:
                        worksheet = workbook[sheet_name]
                        for row in worksheet.rows:
                            for cell in row:
                                if cell.value:
                                    cell.font = Font(name='微软雅黑', size=10)
                    
                # 保存文件
                    writer.close()
                    return True, f"Excel文档已生成: {filename}"
                
                except ImportError:
                    return False, "缺少pandas或openpyxl库，请安装: pip install pandas openpyxl"
                except Exception as e:
                    return False, f"生成Excel文档时出错: {str(e)}"
            
            else:
                return False, "不支持的输出格式"
                
        except Exception as e:
            return False, f"生成文档时出错: {str(e)}"
    
    def _generate_header_doc(self, document, protocol_key=None):
        """生成协议头部分文档"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        heading = document.add_heading('1. 协议头定义', level=1)
        # 设置标题字体
        for run in heading.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        para = document.add_paragraph('协议头定义了所有协议共用的起始字段结构')
        # 设置段落字体
        for run in para.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 添加协议头表格
        if protocol_key and protocol_key in self.protocols:
            protocols = [self.protocols[protocol_key]]
        else:
            # 如果没有指定协议，只过滤type为protocol的协议
            protocols = [p for p in self.protocols.values() if p.get('type') == 'protocol']
        
        if protocols:
            protocol = protocols[0]
            header_fields = protocol.get('header_fields', [])
            
            if header_fields:
                table = document.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # 设置表头
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '字段名称'
                hdr_cells[1].text = '字节类型'
                hdr_cells[2].text = '字段说明'
                
                # 设置表头字体
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '微软雅黑'
                            run.font.bold = True
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                # 填充数据
                for field in header_fields:
                    row_cells = table.add_row().cells
                    row_cells[0].text = field.get('name', '')
                    byte_count = field.get('end_pos', 0) - field.get('start_pos', 0) + 1
                    row_cells[1].text = f"{field.get('type', '')}({byte_count})"
                    row_cells[2].text = field.get('description', '')
                    
                    # 设置数据单元格字体
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '微软雅黑'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                para = document.add_paragraph('未定义协议头字段')
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            para = document.add_paragraph('未找到有效的协议定义')
            for run in para.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _generate_protocol_list_doc(self, document, protocol_key=None):
        """生成协议号列表文档"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        heading = document.add_heading('2. 协议号列表', level=1)
        # 设置标题字体
        for run in heading.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        para = document.add_paragraph('按照协议号从小到大排序')
        # 设置段落字体
        for run in para.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 筛选协议
        if protocol_key and protocol_key in self.protocols:
            protocol = self.protocols[protocol_key]
            if protocol.get('type') == 'protocol':
                protocols = [protocol]
        else:
                # 如果指定了命令，找到其所属的协议
                protocol_name = protocol.get('protocol_name', '')
                protocols = [p for p in self.protocols.values() 
                          if p.get('type') == 'protocol' and p.get('name') == protocol_name]
        else:
            # 只选择类型为protocol的协议
            protocols = [p for p in self.protocols.values() if p.get('type') == 'protocol']
        
        if protocols:
            # 按协议号排序
            try:
            sorted_protocols = sorted(protocols, key=lambda p: int(p.get('protocol_id_dec', '0') or '0'))
            except (ValueError, TypeError):
                # 如果排序失败，使用原始顺序
                sorted_protocols = protocols
            
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # 设置表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '协议号'
            hdr_cells[1].text = '协议说明'
            
            # 设置表头字体
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run.font.bold = True
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 填充数据
            for protocol in sorted_protocols:
                row_cells = table.add_row().cells
                protocol_id = protocol.get('protocol_id_hex', '')
                protocol_id_dec = protocol.get('protocol_id_dec', '')
                row_cells[0].text = f"0x{protocol_id} ({protocol_id_dec})"
                row_cells[1].text = f"{protocol.get('name', '')} - {protocol.get('description', '')}"
                
                # 设置数据单元格字体
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '微软雅黑'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            para = document.add_paragraph('未找到有效的协议定义')
            for run in para.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _generate_protocol_fields_doc(self, document, protocol_key=None):
        """生成协议字段详情文档"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        heading = document.add_heading('3. 协议字段详情', level=1)
        # 设置标题字体
        for run in heading.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 筛选协议
        if protocol_key and protocol_key in self.protocols:
            specified_protocol = self.protocols[protocol_key]
            
            if specified_protocol.get('type') == 'protocol':
                # 如果指定的是一个协议，获取此协议及其所有命令
                protocol_name = specified_protocol.get('name', '')
                protocol_commands = self.get_protocol_commands(protocol_name)
                
                # 只包含指定的协议
                protocols = [specified_protocol]
                
                # 添加此协议的所有命令
                protocols.extend(protocol_commands) if protocol_commands else None
        else:
                # 如果指定的是一个命令，只显示该命令
                protocols = [specified_protocol]
        else:
            # 如果没有指定协议，显示所有协议和命令
            protocols = list(self.protocols.values())
        
        # 按类型分组：先显示所有协议，再显示所有命令
        protocol_types = {'protocol': [], 'command': []}
        
            for protocol in protocols:
            protocol_type = protocol.get('type', 'unknown')
            if protocol_type in protocol_types:
                protocol_types[protocol_type].append(protocol)
        
        # 处理所有协议
        if protocol_types['protocol']:
            try:
                # 按协议ID排序
                protocol_types['protocol'].sort(key=lambda p: int(p.get('protocol_id_dec', '0') or '0'))
            except (ValueError, TypeError):
                pass  # 如果排序失败，使用原始顺序
            
            for i, protocol in enumerate(protocol_types['protocol']):
                protocol_name = protocol.get('name', '未命名协议')
                protocol_id = protocol.get('protocol_id_hex', '')
                protocol_id_dec = protocol.get('protocol_id_dec', '')
                
                # 添加协议标题
                sub_heading = document.add_heading(f"3.{i+1} {protocol_name} (0x{protocol_id})", level=2)
                # 设置子标题字体
                for run in sub_heading.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                para = document.add_paragraph(protocol.get('description', '无描述'))
                # 设置段落字体
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                # 添加字段表格
                fields = protocol.get('fields', [])
                if fields:
                    table = document.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    
                    # 设置表头
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '字段名称'
                    hdr_cells[1].text = '字节类型'
                    hdr_cells[2].text = '字段说明'
                    
                    # 设置表头字体
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '微软雅黑'
                                run.font.bold = True
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    
                    # 填充数据
                    for field in fields:
                        row_cells = table.add_row().cells
                        row_cells[0].text = field.get('name', '??')
                        byte_count = field.get('end_pos', 0) - field.get('start_pos', 0) + 1
                        row_cells[1].text = f"{field.get('type', 'unknown')}({byte_count})"
                        row_cells[2].text = field.get('description', '')
                        
                        # 设置数据单元格字体
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = '微软雅黑'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                else:
                    para = document.add_paragraph('此协议未定义字段')
                    # 设置字体
                    for run in para.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                # 添加分隔符
                document.add_paragraph()
        
        # 处理所有命令
        if protocol_types['command']:
            cmd_heading = document.add_heading('4. 命令字段详情', level=1)
            # 设置标题字体
            for run in cmd_heading.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            try:
                # 按协议ID排序
                protocol_types['command'].sort(key=lambda p: int(p.get('protocol_id_dec', '0') or '0'))
            except (ValueError, TypeError):
                pass  # 如果排序失败，使用原始顺序
            
            for i, command in enumerate(protocol_types['command']):
                command_name = command.get('name', '未命名命令')
                command_id = command.get('protocol_id_hex', '')
                command_id_dec = command.get('protocol_id_dec', '')
                protocol_name = command.get('protocol_name', '')
                
                # 添加命令标题
                cmd_sub_heading = document.add_heading(f"4.{i+1} {command_name} (0x{command_id})", level=2)
                # 设置子标题字体
                for run in cmd_sub_heading.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                para = document.add_paragraph(f"协议: {protocol_name} - {command.get('description', '无描述')}")
                # 设置段落字体
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                # 添加字段表格
                fields = command.get('fields', [])
                if fields:
                    table = document.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    
                    # 设置表头
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '字段名称'
                    hdr_cells[1].text = '字节类型'
                    hdr_cells[2].text = '字段说明'
                    
                    # 设置表头字体
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '微软雅黑'
                                run.font.bold = True
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    
                    # 填充数据
                    for field in fields:
                        row_cells = table.add_row().cells
                        row_cells[0].text = field.get('name', '??')
                        byte_count = field.get('end_pos', 0) - field.get('start_pos', 0) + 1
                        row_cells[1].text = f"{field.get('type', 'unknown')}({byte_count})"
                        row_cells[2].text = field.get('description', '')
                        
                        # 设置数据单元格字体
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = '微软雅黑'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
                    para = document.add_paragraph('此命令未定义字段')
                    # 设置字体
                    for run in para.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                # 添加分隔符
                document.add_paragraph()
        
        # 如果没有协议和命令
        if not protocol_types['protocol'] and not protocol_types['command']:
            para = document.add_paragraph('未找到有效的协议或命令定义')
            # 设置字体
            for run in para.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _generate_header_excel(self, writer, protocol_key=None):
        """生成Excel格式的协议头文档"""
        import pandas as pd
        
        # 筛选协议
        if protocol_key and protocol_key in self.protocols:
            protocols = [self.protocols[protocol_key]]
        else:
            # 如果没有指定协议，只选择类型为protocol的协议
            protocols = [p for p in self.protocols.values() if p.get('type') == 'protocol']
        
        if protocols:
            protocol = protocols[0]
            header_fields = protocol.get('header_fields', [])
            
            if header_fields:
                # 准备数据
                data = []
                for field in header_fields:
                    byte_count = field.get('end_pos', 0) - field.get('start_pos', 0) + 1
                    data.append({
                        '字段名称': field.get('name', ''),
                        '字节类型': f"{field.get('type', '')}({byte_count})",
                        '字段说明': field.get('description', '')
                    })
                
                # 创建DataFrame并写入Excel
                df = pd.DataFrame(data)
                df.to_excel(writer, sheet_name='协议头定义', index=False)
                
                # 设置列宽和样式
                worksheet = writer.sheets['协议头定义']
                worksheet.column_dimensions['A'].width = 20
                worksheet.column_dimensions['B'].width = 15
                worksheet.column_dimensions['C'].width = 40
                
                # 设置表头样式
                from openpyxl.styles import Font, Alignment, PatternFill
                header_font = Font(name='微软雅黑', bold=True, size=11)
                header_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
                
                for cell in worksheet[1]:
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center', vertical='center')
            else:
                # 创建空表格
                df = pd.DataFrame({'提示': ['未定义协议头字段']})
                df.to_excel(writer, sheet_name='协议头定义', index=False)
        else:
            # 创建空表格
            df = pd.DataFrame({'提示': ['未找到有效的协议定义']})
            df.to_excel(writer, sheet_name='协议头定义', index=False)
    
    def _generate_protocol_list_excel(self, writer, protocol_key=None):
        """生成Excel格式的协议列表"""
        import pandas as pd
        
        # 筛选协议
        if protocol_key and protocol_key in self.protocols:
            protocol = self.protocols[protocol_key]
            if protocol.get('type') == 'protocol':
                protocols = [protocol]
        else:
                # 如果指定了命令，找到其所属的协议
                protocol_name = protocol.get('protocol_name', '')
                protocols = [p for p in self.protocols.values() 
                          if p.get('type') == 'protocol' and p.get('name') == protocol_name]
        else:
            # 只选择类型为protocol的协议
            protocols = [p for p in self.protocols.values() if p.get('type') == 'protocol']
        
        if protocols:
            # 准备数据
            data = []
            for protocol in protocols:
                data.append({
                    '协议号': f"0x{protocol.get('protocol_id_hex', '')} ({protocol.get('protocol_id_dec', '')})",
                    '协议说明': f"{protocol.get('name', '')} - {protocol.get('description', '')}"
                })
            
            # 创建DataFrame并写入Excel
            df = pd.DataFrame(data)
            df.to_excel(writer, sheet_name='协议号列表', index=False)
            
            # 设置列宽和样式
            worksheet = writer.sheets['协议号列表']
            worksheet.column_dimensions['A'].width = 15
            worksheet.column_dimensions['B'].width = 40
            
            # 设置表头样式
            from openpyxl.styles import Font, Alignment, PatternFill
            header_font = Font(name='微软雅黑', bold=True, size=11)
            header_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
            
            for cell in worksheet[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center')
        else:
            # 创建空表格
            df = pd.DataFrame({'提示': ['未找到有效的协议定义']})
            df.to_excel(writer, sheet_name='协议号列表', index=False)
    
    def _generate_protocol_fields_excel(self, writer, protocol_key=None):
        """生成Excel格式的协议字段详情"""
        import pandas as pd
        
        # 筛选协议
        if protocol_key and protocol_key in self.protocols:
            specified_protocol = self.protocols[protocol_key]
            
            if specified_protocol.get('type') == 'protocol':
                # 如果指定的是一个协议，获取此协议及其所有命令
                protocol_name = specified_protocol.get('name', '')
                protocol_commands = self.get_protocol_commands(protocol_name)
                
                # 只包含指定的协议
                protocols = [specified_protocol]
                
                # 添加此协议的所有命令
                protocols.extend(protocol_commands) if protocol_commands else None
        else:
                # 如果指定的是一个命令，只显示该命令
                protocols = [specified_protocol]
        else:
            # 如果没有指定协议，显示所有协议和命令
            protocols = list(self.protocols.values())
        
        # 按类型分组：先显示所有协议，再显示所有命令
        protocol_types = {'protocol': [], 'command': []}
        
            for protocol in protocols:
            protocol_type = protocol.get('type', 'unknown')
            if protocol_type in protocol_types:
                protocol_types[protocol_type].append(protocol)
        
        # 设置样式相关
        from openpyxl.styles import Font, Alignment, PatternFill
        header_font = Font(name='微软雅黑', bold=True, size=11)
        header_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
        
        # 处理所有协议
        if protocol_types['protocol']:
            # 准备协议数据
            all_protocol_data = []
            
            for protocol in protocol_types['protocol']:
                protocol_name = protocol.get('name', '未命名协议')
                protocol_id = protocol.get('protocol_id_hex', '')
                protocol_id_dec = protocol.get('protocol_id_dec', '')
                
                fields = protocol.get('fields', [])
                if fields:
                    for field in fields:
                        byte_count = field.get('end_pos', 0) - field.get('start_pos', 0) + 1
                        all_protocol_data.append({
                            '协议': f"{protocol_name} (0x{protocol_id})",
                            '字段名称': field.get('name', '??'),
                            '字节类型': f"{field.get('type', 'unknown')}({byte_count})",
                            '字段说明': field.get('description', '')
                        })
                else:
                    all_protocol_data.append({
                        '协议': f"{protocol_name} (0x{protocol_id})",
                        '字段名称': '无字段',
                        '字节类型': '',
                        '字段说明': ''
                        })
                    
                    # 创建DataFrame并写入Excel
            if all_protocol_data:
                df = pd.DataFrame(all_protocol_data)
                df.to_excel(writer, sheet_name='协议字段详情', index=False)
                
                # 设置列宽和样式
                worksheet = writer.sheets['协议字段详情']
                worksheet.column_dimensions['A'].width = 25
                worksheet.column_dimensions['B'].width = 20
                worksheet.column_dimensions['C'].width = 15
                worksheet.column_dimensions['D'].width = 40
                
                # 设置表头样式
                for cell in worksheet[1]:
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                else:
                    # 创建空表格
                df = pd.DataFrame({'提示': ['协议未定义字段']})
                df.to_excel(writer, sheet_name='协议字段详情', index=False)
        
        # 处理所有命令
        if protocol_types['command']:
            # 准备命令数据
            all_command_data = []
            
            for command in protocol_types['command']:
                command_name = command.get('name', '未命名命令')
                command_id = command.get('protocol_id_hex', '')
                protocol_name = command.get('protocol_name', '')
                
                fields = command.get('fields', [])
                if fields:
                    for field in fields:
                        byte_count = field.get('end_pos', 0) - field.get('start_pos', 0) + 1
                        all_command_data.append({
                            '命令': f"{command_name} (0x{command_id})",
                            '所属协议': protocol_name,
                            '字段名称': field.get('name', '??'),
                            '字节类型': f"{field.get('type', 'unknown')}({byte_count})",
                            '字段说明': field.get('description', '')
                        })
                else:
                    all_command_data.append({
                        '命令': f"{command_name} (0x{command_id})",
                        '所属协议': protocol_name,
                        '字段名称': '无字段',
                        '字节类型': '',
                        '字段说明': ''
                    })
            
            # 创建DataFrame并写入Excel
            if all_command_data:
                df = pd.DataFrame(all_command_data)
                df.to_excel(writer, sheet_name='命令字段详情', index=False)
                
                # 设置列宽和样式
                worksheet = writer.sheets['命令字段详情']
                worksheet.column_dimensions['A'].width = 25
                worksheet.column_dimensions['B'].width = 20
                worksheet.column_dimensions['C'].width = 20
                worksheet.column_dimensions['D'].width = 15
                worksheet.column_dimensions['E'].width = 40
                
                # 设置表头样式
                for cell in worksheet[1]:
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center', vertical='center')
        else:
            # 创建空表格
                df = pd.DataFrame({'提示': ['命令未定义字段']})
                df.to_excel(writer, sheet_name='命令字段详情', index=False)

    def save_command(self, group, command_id, command_data):
        """保存命令数据到文件"""
        # 确保命令数据包含必要的字段
        command_data["protocol_id_hex"] = command_id
        command_data["protocol_id"] = command_id
        command_data["group"] = group
        command_data["type"] = "command"
        
        try:
            # 计算十进制值
            command_data["protocol_id_dec"] = str(int(command_id, 16))
        except ValueError:
            command_data["protocol_id_dec"] = "未知"
        
        # 创建协议命令目录
        protocol_dir = self.data_dir / group
        protocol_dir.mkdir(exist_ok=True, parents=True)
        
        # 命令文件以命令ID命名
        file_path = protocol_dir / f"{command_id}.json"
        
        try:
            print(f"保存命令到文件: {file_path}")
            
            # 如果文件已存在，读取现有命令
            existing_commands = {}
            if file_path.exists():
                with open(file_path, 'r', encoding='utf-8') as f:
                    try:
                        file_content = json.load(f)
                        # 处理不同格式的文件内容
                        if isinstance(file_content, list):
                            # 如果是列表，转换为新格式
                            existing_commands = {command_id: file_content}
                        elif isinstance(file_content, dict):
                            if 'name' in file_content:
                                # 如果是单个命令对象，转换为新格式
                                existing_commands = {command_id: [file_content]}
                            else:
                                # 如果是新格式，直接使用
                                existing_commands = file_content
                    except json.JSONDecodeError:
                        existing_commands = {command_id: []}
            
            # 初始化该命令ID的命令列表
            if command_id not in existing_commands:
                existing_commands[command_id] = []
            
            # 检查是否已存在相同名称的命令
            command_name = command_data.get("name", "unknown")
            command_exists = False
            for i, cmd in enumerate(existing_commands[command_id]):
                if isinstance(cmd, dict) and cmd.get("name") == command_name:
                    # 如果存在相同名称的命令，更新它
                    existing_commands[command_id][i] = command_data
                    command_exists = True
                    break
            
            if not command_exists:
                # 如果不存在相同名称的命令，添加新命令
                existing_commands[command_id].append(command_data)
            
            # 保存所有命令到文件
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(existing_commands, f, ensure_ascii=False, indent=2)
            
            # 更新内存中的命令数据
            full_key = f"{group}/{command_id}"
            self.protocols[full_key] = existing_commands[command_id]
            
            # 更新命令字典
            protocol_name = command_data.get("protocol_name", "")
            if protocol_name:
                if protocol_name not in self.protocol_commands:
                    self.protocol_commands[protocol_name] = {}
                self.protocol_commands[protocol_name][command_id] = existing_commands[command_id]
            
            print(f"命令保存成功: {full_key}")
            return True, f"命令已保存: {command_id} (十进制: {command_data.get('protocol_id_dec', '未知')}) 到 {group}"
        except Exception as e:
            print(f"保存命令失败: {e}")
            return False, f"保存命令失败: {e}"

    def import_commands(self, file_path):
        """从文件导入命令"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                commands = json.load(f)
                
            # 转换旧格式到新格式
            for protocol_name, protocol_commands in commands.items():
                if protocol_name not in self.protocol_commands:
                    self.protocol_commands[protocol_name] = {}
                    
                for command_id, command_list in protocol_commands.items():
                    if command_id not in self.protocol_commands[protocol_name]:
                        self.protocol_commands[protocol_name][command_id] = []
                        
                    if isinstance(command_list, list):
                        self.protocol_commands[protocol_name][command_id].extend(command_list)
                    else:
                        self.protocol_commands[protocol_name][command_id].append(command_list)
                        
            return True, "命令导入成功"
        except Exception as e:
            return False, f"导入命令失败: {str(e)}"

    def merge_commands(self, source_group, target_group):
        """合并命令组"""
        try:
            if source_group not in self.protocol_commands or target_group not in self.protocol_commands:
                return False, "源组或目标组不存在"
                
            # 合并命令
            for command_id, command_list in self.protocol_commands[source_group].items():
                if command_id not in self.protocol_commands[target_group]:
                    self.protocol_commands[target_group][command_id] = []
                    
                if isinstance(command_list, list):
                    self.protocol_commands[target_group][command_id].extend(command_list)
                else:
                    self.protocol_commands[target_group][command_id].append(command_list)
                    
            return True, "命令合并成功"
        except Exception as e:
            return False, f"合并命令失败: {str(e)}"
